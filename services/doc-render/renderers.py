"""
The three document renderers (WARP-2211): spec in, bytes out.

Every function here is pure — it takes a validated spec and returns bytes. No
filesystem, no network, no credentials. The orchestrator owns auth, path
validation and the Nextcloud upload; this module owns nothing but the file
format. That split is why the container can hold no user token at all.

Libraries, all permissive (shipping the appliance is conveyance, so the
permissive-only rule applies):

    python-docx  MIT            .docx
    openpyxl     MIT            .xlsx
    reportlab    BSD-3-Clause   .pdf   — pure Python, no cairo/pango, so no
                                        LGPL native layer enters the image.
                                        WeasyPrint was rejected for exactly
                                        that; wkhtmltopdf is LGPL outright.
"""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from markdown_blocks import parse_blocks, split_inline, to_rl_markup

MIME = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

# Caps. A legitimate spec is bounded by the model's 4096-token output budget,
# so these only ever bite on a malformed or hostile payload.
MAX_SHEETS = 24
MAX_ROWS_PER_SHEET = 50_000
MAX_COLUMNS = 256
# Excel's own hard ceiling on a cell's contents. Longer values are truncated
# rather than rejected: silently writing a value Excel cannot reopen would
# produce a corrupt-looking workbook, and failing the whole render over one
# long cell is worse than a visibly clipped one.
MAX_CELL_CHARS = 32_767
# Sheet names: Excel forbids these characters and caps the name at 31 chars.
INVALID_SHEET_CHARS = set(r"[]:*?/\\")
# WARP-2521 — formula-injection guard for string cells (CWE-1236). A cell
# whose first non-whitespace character is = + - @ executes as a live formula
# the moment Excel/Sheets/Numbers opens the workbook, and cell values here are
# LLM/user-supplied. Mirrors the dashboard CSV export's FORMULA_LEADER
# (apps/web-dashboard/src/lib/audit-csv.ts, WARP-1031): leading whitespace is
# included because spreadsheet apps trim space/tab/CR before deciding a cell
# is a formula, so "\t=cmd" executes like "=cmd". A leading apostrophe makes
# every one of them render as text.
FORMULA_LEADER = re.compile(r"^\s*[=+\-@]")


class RenderError(ValueError):
    """A spec the renderer refuses. Surfaces to the caller as 400."""


# ── .docx ────────────────────────────────────────────────────────────────


def render_docx(title: str, body_markdown: str) -> bytes:
    doc = Document()
    if title:
        doc.add_heading(title, level=0)

    for block in parse_blocks(body_markdown):
        if block.kind == "heading":
            doc.add_heading(block.text, level=min(block.level, 4))
        elif block.kind == "paragraph":
            _docx_runs(doc.add_paragraph(), block.text)
        elif block.kind == "bullets":
            for item in block.items:
                _docx_runs(doc.add_paragraph(style="List Bullet"), item)
        elif block.kind == "numbers":
            for item in block.items:
                _docx_runs(doc.add_paragraph(style="List Number"), item)
        elif block.kind == "table":
            _docx_table(doc, block.header, block.rows)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_runs(paragraph: Any, text: str) -> None:
    """Add text to a paragraph as bold/italic runs.

    Word has no markup string: a run's boldness is a property of the run
    object, so inline emphasis has to be split into spans and applied here.
    """
    for run_text, bold, italic in split_inline(text):
        run = paragraph.add_run(run_text)
        run.bold = bold
        run.italic = italic


def _docx_table(doc: Any, header: list[str], rows: list[list[str]]) -> None:
    width = max([len(header)] + [len(r) for r in rows]) if (header or rows) else 0
    if width == 0:
        return
    table = doc.add_table(rows=1, cols=width)
    table.style = "Table Grid"
    for i, cell_text in enumerate(header[:width]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(cell_text)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i in range(width):
            # Ragged rows are normal in hand-written Markdown tables; pad
            # rather than raising, so one short row cannot fail the render.
            cells[i].text = row[i] if i < len(row) else ""


# ── .pdf ─────────────────────────────────────────────────────────────────


def render_pdf(title: str, body_markdown: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.9 * inch,
        bottomMargin=0.9 * inch,
        title=title or "Document",
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "DropletBody",
        parent=styles["BodyText"],
        fontSize=10.5,
        leading=15,
        alignment=TA_LEFT,
        spaceAfter=8,
    )
    story: list[Any] = []

    if title:
        story.append(Paragraph(to_rl_markup(title), styles["Title"]))
        story.append(Spacer(1, 10))

    for block in parse_blocks(body_markdown):
        if block.kind == "heading":
            style = styles[f"Heading{min(block.level, 3)}"]
            story.append(Paragraph(to_rl_markup(block.text), style))
        elif block.kind == "paragraph":
            story.append(Paragraph(to_rl_markup(block.text), body))
        elif block.kind in ("bullets", "numbers"):
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(to_rl_markup(i), body)) for i in block.items],
                    bulletType="bullet" if block.kind == "bullets" else "1",
                    leftIndent=18,
                )
            )
            story.append(Spacer(1, 6))
        elif block.kind == "table":
            flowable = _pdf_table(block.header, block.rows, body)
            if flowable is not None:
                story.append(flowable)
                story.append(Spacer(1, 10))

    # A PDF with no pages is invalid. An empty body is a legitimate request
    # (a titled cover sheet), so emit a blank paragraph rather than refusing.
    if not story:
        story.append(Paragraph("", body))

    doc.build(story)
    return buf.getvalue()


def _pdf_table(header: list[str], rows: list[list[str]], body: Any) -> Any | None:
    width = max([len(header)] + [len(r) for r in rows]) if (header or rows) else 0
    if width == 0:
        return None
    data: list[list[Any]] = []
    if header:
        data.append([Paragraph(f"<b>{to_rl_markup(c)}</b>", body) for c in _pad(header, width)])
    for row in rows:
        data.append([Paragraph(to_rl_markup(c), body) for c in _pad(row, width)])
    table = Table(data, repeatRows=1 if header else 0, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c9ccd4")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef0f4")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def _pad(row: list[str], width: int) -> list[str]:
    return list(row[:width]) + [""] * max(0, width - len(row))


# ── .xlsx ────────────────────────────────────────────────────────────────


def render_xlsx(sheets: list[dict[str, Any]]) -> bytes:
    if not sheets:
        raise RenderError("at least one sheet is required")
    if len(sheets) > MAX_SHEETS:
        raise RenderError(f"too many sheets (max {MAX_SHEETS})")

    wb = Workbook()
    wb.remove(wb.active)
    used: set[str] = set()

    for index, spec in enumerate(sheets):
        if not isinstance(spec, dict):
            raise RenderError("each sheet must be an object")
        columns = spec.get("columns") or []
        rows = spec.get("rows") or []
        if not isinstance(columns, list) or not isinstance(rows, list):
            raise RenderError("sheet columns and rows must be arrays")
        if len(rows) > MAX_ROWS_PER_SHEET:
            raise RenderError(f"too many rows (max {MAX_ROWS_PER_SHEET})")
        if len(columns) > MAX_COLUMNS:
            raise RenderError(f"too many columns (max {MAX_COLUMNS})")

        ws = wb.create_sheet(_sheet_name(spec.get("name"), index, used))

        if columns:
            ws.append([_cell(c) for c in columns[:MAX_COLUMNS]])
            for cell in ws[1]:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(vertical="top")
            ws.freeze_panes = "A2"

        for row in rows:
            if not isinstance(row, list):
                raise RenderError("each row must be an array")
            ws.append([_cell(v) for v in row[:MAX_COLUMNS]])

        _autofit(ws, columns, rows)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _sheet_name(raw: Any, index: int, used: set[str]) -> str:
    """Coerce a requested sheet name into one Excel will accept.

    Excel rejects []:*?/\\ and caps names at 31 characters, and a duplicate
    name raises inside openpyxl. Sanitising beats refusing: the model picked
    a human title, and losing a colon is a better outcome than failing the
    whole workbook over it.
    """
    name = str(raw).strip() if raw not in (None, "") else f"Sheet{index + 1}"
    name = "".join(ch for ch in name if ch not in INVALID_SHEET_CHARS)[:31].strip()
    if not name:
        name = f"Sheet{index + 1}"
    base, n = name, 2
    while name.lower() in used:
        suffix = f" ({n})"
        name = base[: 31 - len(suffix)] + suffix
        n += 1
    used.add(name.lower())
    return name


def _cell(value: Any) -> Any:
    """Pass native scalars through; stringify everything else.

    Numbers and booleans reach the sheet as numbers and booleans so they can
    be summed and filtered. Nothing is COERCED, though: a numeric-looking
    string stays a string, because guessing turns a zip code or a phone
    number into a number and silently loses the leading zero. Same reasoning
    as `convert_data_format`'s refusal to infer CSV cell types.

    Strings are additionally screened for formula leaders (WARP-2521, see
    FORMULA_LEADER): openpyxl writes a string starting with `=` as a live
    formula, so an LLM/user-supplied `=HYPERLINK(...)` would execute on open.
    The apostrophe goes on BEFORE the MAX_CELL_CHARS truncate so the guard
    can never push a cell past Excel's own hard ceiling.
    """
    if value is None:
        return ""
    if isinstance(value, bool) or isinstance(value, (int, float)):
        return value
    text = str(value)
    if FORMULA_LEADER.match(text):
        text = "'" + text
    return text[:MAX_CELL_CHARS]


def _autofit(ws: Any, columns: list[Any], rows: list[Any]) -> None:
    """Width each column to its widest cell, within reason.

    openpyxl has no real autofit (that is a rendering-time measurement Excel
    does), so this approximates from character counts — enough that a report
    opens readable instead of showing a column of `####`.
    """
    widths: dict[int, int] = {}
    for i, col in enumerate(columns[:MAX_COLUMNS], start=1):
        widths[i] = len(str(col))
    for row in rows[:500]:
        if not isinstance(row, list):
            continue
        for i, value in enumerate(row[:MAX_COLUMNS], start=1):
            widths[i] = max(widths.get(i, 0), len(str(value if value is not None else "")))
    for i, width in widths.items():
        ws.column_dimensions[get_column_letter(i)].width = min(max(width + 2, 8), 60)
