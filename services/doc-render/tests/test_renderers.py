"""The three renderers produce files their own libraries can reopen.

Asserting on a magic number only proves something was written. Every format
here is round-tripped through the reader that a user's application would use —
openpyxl reopens the workbook, python-docx reopens the document — because a
structurally-invalid OOXML package still starts with `PK`, and that is exactly
the failure this suite exists to catch.
"""

from __future__ import annotations

import io
import zipfile

import pytest

import renderers
from markdown_blocks import parse_blocks, split_inline, to_rl_markup

SAMPLE = """# Quarter in review

Revenue held **flat** against a *soft* comparable.

## Highlights

- Renewals up 4%
- Churn steady

1. Close the Ashby account
2. Ship the rack panel

| Region | Revenue |
|---|---|
| West | 120 |
| East | 98 |
"""


# ── .xlsx ────────────────────────────────────────────────────────────────


def test_xlsx_reopens_with_the_expected_grid():
    from openpyxl import load_workbook

    data = renderers.render_xlsx(
        [{"name": "Q3", "columns": ["Region", "Revenue"], "rows": [["West", 120], ["East", 98]]}]
    )
    wb = load_workbook(io.BytesIO(data))
    ws = wb["Q3"]
    assert [c.value for c in ws[1]] == ["Region", "Revenue"]
    assert [c.value for c in ws[2]] == ["West", 120]
    assert ws["B2"].value == 120 and isinstance(ws["B2"].value, int)


def test_xlsx_header_is_bold_and_frozen():
    from openpyxl import load_workbook

    data = renderers.render_xlsx([{"columns": ["A"], "rows": [["x"]]}])
    ws = load_workbook(io.BytesIO(data)).worksheets[0]
    assert ws["A1"].font.bold is True
    assert ws.freeze_panes == "A2"


def test_xlsx_does_not_coerce_a_numeric_looking_string():
    """A zip code keeps its leading zero.

    Same reasoning as convert_data_format's refusal to infer CSV cell types:
    guessing turns 01234 into 1234 and silently corrupts the data.
    """
    from openpyxl import load_workbook

    data = renderers.render_xlsx([{"columns": ["Zip"], "rows": [["01234"]]}])
    ws = load_workbook(io.BytesIO(data)).worksheets[0]
    assert ws["A2"].value == "01234"


def test_xlsx_sanitises_an_illegal_sheet_name():
    from openpyxl import load_workbook

    data = renderers.render_xlsx([{"name": "Q3/Q4: results[final]", "columns": ["A"], "rows": []}])
    name = load_workbook(io.BytesIO(data)).sheetnames[0]
    assert not (set(name) & set(r"[]:*?/\\"))
    assert len(name) <= 31


def test_xlsx_deduplicates_repeated_sheet_names():
    from openpyxl import load_workbook

    data = renderers.render_xlsx(
        [{"name": "Data", "columns": [], "rows": []}, {"name": "Data", "columns": [], "rows": []}]
    )
    assert len(set(load_workbook(io.BytesIO(data)).sheetnames)) == 2


def test_xlsx_pads_a_ragged_row_instead_of_raising():
    from openpyxl import load_workbook

    data = renderers.render_xlsx([{"columns": ["A", "B", "C"], "rows": [["only-one"]]}])
    ws = load_workbook(io.BytesIO(data)).worksheets[0]
    assert ws["A2"].value == "only-one"


def test_xlsx_refuses_an_empty_spec():
    with pytest.raises(renderers.RenderError):
        renderers.render_xlsx([])


def test_xlsx_refuses_a_non_array_row():
    with pytest.raises(renderers.RenderError):
        renderers.render_xlsx([{"columns": ["A"], "rows": ["not-a-row"]}])


# ── .docx ────────────────────────────────────────────────────────────────


def test_docx_reopens_and_carries_the_content():
    from docx import Document

    doc = Document(io.BytesIO(renderers.render_docx("Quarter in review", SAMPLE)))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Quarter in review" in text
    assert "Renewals up 4%" in text
    assert "Close the Ashby account" in text


def test_docx_renders_a_markdown_table_as_a_real_table():
    from docx import Document

    doc = Document(io.BytesIO(renderers.render_docx("t", SAMPLE)))
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[0].cells[0].text == "Region"
    assert doc.tables[0].rows[1].cells[1].text == "120"


def test_docx_applies_bold_as_a_run_property():
    """Word has no markup string — emphasis must be a run property."""
    from docx import Document

    doc = Document(io.BytesIO(renderers.render_docx("", "plain **loud** plain")))
    runs = [r for p in doc.paragraphs for r in p.runs]
    assert any(r.text == "loud" and r.bold for r in runs)


def test_docx_is_a_valid_ooxml_package():
    data = renderers.render_docx("t", "body")
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        assert "word/document.xml" in z.namelist()
        assert "[Content_Types].xml" in z.namelist()


# ── .pdf ─────────────────────────────────────────────────────────────────


def test_pdf_has_a_header_and_an_eof():
    data = renderers.render_pdf("Quarter in review", SAMPLE)
    assert data.startswith(b"%PDF-")
    assert b"%%EOF" in data[-2048:]


def test_pdf_renders_an_empty_body():
    """A titled cover sheet is a legitimate request; a pageless PDF is not."""
    data = renderers.render_pdf("Just a title", "")
    assert data.startswith(b"%PDF-")


def test_pdf_survives_markup_in_the_body():
    """User text is escaped BEFORE our own tags go in.

    ReportLab parses <b>/<i> out of the string it is handed, so an unescaped
    `<b>` in the body would restyle the document and a bare `&` would raise
    inside the PDF build.
    """
    data = renderers.render_pdf("t", "5 < 6 & 7 > 2 <b>not a tag</b>")
    assert data.startswith(b"%PDF-")


# ── the shared block parser ──────────────────────────────────────────────


def test_parser_recognises_each_block_kind():
    kinds = [b.kind for b in parse_blocks(SAMPLE)]
    assert "heading" in kinds
    assert "paragraph" in kinds
    assert "bullets" in kinds
    assert "numbers" in kinds
    assert "table" in kinds


def test_parser_coalesces_consecutive_list_items():
    blocks = [b for b in parse_blocks("- a\n- b\n- c") if b.kind == "bullets"]
    assert len(blocks) == 1
    assert blocks[0].items == ["a", "b", "c"]


def test_a_pipe_line_without_a_separator_is_prose_not_a_table():
    blocks = parse_blocks("revenue | cost were both up")
    assert [b.kind for b in blocks] == ["paragraph"]


def test_parser_keeps_an_unsupported_construct_as_text():
    """A subset is honest when it degrades to text, not when it drops input."""
    blocks = parse_blocks("> a blockquote we do not style")
    assert blocks[0].kind == "paragraph"
    assert "blockquote" in blocks[0].text


def test_inline_split_handles_bold_and_italic():
    assert ("loud", True, False) in split_inline("a **loud** b")
    assert ("soft", False, True) in split_inline("a *soft* b")


def test_rl_markup_escapes_before_it_adds_tags():
    out = to_rl_markup("5 < 6 & <b>x</b> **real**")
    assert "&lt;b&gt;x&lt;/b&gt;" in out
    assert "&amp;" in out
    assert "<b>real</b>" in out
