"""DOCX extractor using python-docx.

Flattens paragraphs and table rows into document body text. Tables are
joined with ``" | "`` between cells so column structure survives the
chunker.

WARP-287: DOCX is non-MVP for positional anchors — python-docx doesn't
track page breaks reliably, so every Span carries ``NoneAnchor`` (no
deep-link target).

WARP-435 / ADR-003 Phase 1: walks the paragraph stream and tracks
heading-style depth via ``paragraph.style.name`` (``Heading 1``,
``Heading 2``, ...). The unified design preserves that breadcrumb by
emitting one Span *per heading-delimited section*, each carrying the
section's ``section_path`` (the live heading stack). Consecutive
paragraphs that share a section path are merged into a single Span so
the chunker still sees coherent blocks. Documents with no headings emit
a single Span with ``[filename]`` as its section path.
"""
from __future__ import annotations

import os

from docx import Document

from anchor_schema import NoneAnchor
from extractors.spans import Span
from extractors.types import ExtractedDoc


def _heading_level(style_name: str) -> int | None:
    """Parse a Word style name into a heading depth (1..9) or None.

    Word ships ``Heading 1`` .. ``Heading 9`` as the canonical heading
    styles; some templates add custom ``Heading X - Foo`` variants. We
    accept anything that starts with ``Heading `` followed by a number.
    """
    if not style_name or not style_name.startswith("Heading"):
        return None
    # Strip the "Heading" prefix and grab the first space-separated token.
    rest = style_name[len("Heading") :].strip()
    if not rest:
        return None
    first = rest.split()[0]
    try:
        level = int(first)
    except ValueError:
        return None
    if 1 <= level <= 9:
        return level
    return None


def extract(path: str) -> ExtractedDoc:
    document = Document(path)
    filename = os.path.basename(path)

    # We accumulate (section_path, [paragraph, ...]) sections. A new
    # section opens whenever the live heading stack changes; consecutive
    # paragraphs under the same heading path coalesce into one section so
    # the chunker sees a coherent block per Span.
    sections: list[tuple[list[str], list[str]]] = []
    stack: list[str] = []

    def _current_path() -> list[str]:
        # Drop empty padding entries so the breadcrumb stays clean.
        path = [s for s in stack if s]
        return path if path else [filename]

    def _append(text: str) -> None:
        path = _current_path()
        if sections and sections[-1][0] == path:
            sections[-1][1].append(text)
        else:
            sections.append((path, [text]))

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = ""
        try:
            style_name = (para.style.name or "").strip() if para.style else ""
        except Exception:
            style_name = ""

        level = _heading_level(style_name)
        if level is not None:
            # Truncate stack to ``level - 1`` (heading-1 resets to empty,
            # heading-2 keeps the heading-1 above it, etc.) then push the
            # new heading. Pad with empty strings rather than over-popping.
            stack = stack[: level - 1]
            while len(stack) < level - 1:
                stack.append("")
            stack.append(text)

        # The heading paragraph is part of its own section's text.
        _append(text)

    # Tables go too — flatten cell text. Tables inherit whatever section
    # was active when iteration ends; python-docx doesn't expose
    # heading-aware ordering for tables vs paragraphs in a single
    # iteration, so we attach them to the trailing section path.
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                _append(row_text)

    metadata = {
        "extractor_name": "docx",
        "extractor_version": "2",
        "paragraph_count": len(document.paragraphs),
        "word_count": sum(
            len(" ".join(paras).split()) for _, paras in sections
        ),
    }

    spans: list[Span] = []
    for path, paras in sections:
        body = "\n\n".join(paras).strip()
        if not body:
            continue
        spans.append(Span(text=body, anchor=NoneAnchor(), section_path=path))

    return ExtractedDoc(
        spans=spans,
        language=None,
        metadata=metadata,
        warnings=[],
    )
